from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_report():
    doc = Document()
    
    # Title Page / Header
    title = doc.add_heading('NEXATTEND PROJECT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Project Title: NexAttend – AI-Powered Face Recognition Attendance System', style='Subtitle')
    doc.add_paragraph('Technology Stack: Python 3, OpenCV, Tkinter, Pandas, Pillow, Pyttsx3')
    
    # Section 1
    doc.add_heading('1. ABSTRACT', level=1)
    p = doc.add_paragraph()
    p.add_run('NexAttend is an advanced, automated face recognition attendance system designed to replace traditional manual roll-call methods. By leveraging computer vision and machine learning, the system accurately identifies students in real-time and logs their attendance into structured CSV and Excel files. The project features a premium, state-of-the-art UI/UX design, ensuring both aesthetic appeal and high functional performance for institutional use.')
    
    # Section 2
    doc.add_heading('2. INTRODUCTION & OBJECTIVE', level=1)
    p = doc.add_paragraph()
    p.add_run('The primary objective of NexAttend is to streamline the attendance tracking process by integrating real-time facial recognition. Traditional methods are prone to proxy attendance and require significant manual effort. This system automates the process by implementing established computer vision pipelines using Python and OpenCV, drastically reducing administrative overhead and increasing accuracy.')
    
    # Section 3
    doc.add_heading('3. MODELS & ALGORITHMS USED', level=1)
    p = doc.add_paragraph('The core computer vision pipeline is built upon two highly efficient models:')
    
    p1 = doc.add_paragraph(style='List Bullet')
    r1 = p1.add_run('Viola-Jones Algorithm (Haar Cascade V2): ')
    r1.bold = True
    p1.add_run('Used for robust real-time face detection. The system focuses on specific facial regions (eyes and nose) to draw bounding boxes. It utilizes custom visual extraction overlays to dynamically demonstrate the Haar Cascade scanning process directly on the camera HUD.')
    
    p2 = doc.add_paragraph(style='List Bullet')
    r2 = p2.add_run('Local Binary Patterns Histograms (LBPH) Recognizer: ')
    r2.bold = True
    p2.add_run('Used for facial recognition and classification. LBPH is highly resilient under varying lighting conditions and is trained on 100 image samples per student to ensure high confidence matching. Internal confidence scores are translated into intuitive percentage metrics on the UI.')
    
    # Section 4
    doc.add_heading('4. UI/UX DESIGN: VISIOMARK PREMIUM DARK THEME', level=1)
    p = doc.add_paragraph('NexAttend sets itself apart with an exceptional user interface, transforming standard Tkinter into a "VisioMark Premium Dark Theme" that offers a deep space cyber-glass aesthetic:')
    
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run('Color Palette & Glassmorphism: ').bold = True
    p1.add_run('Uses an ultra-deep space blue/black (#050511) background, slightly elevated card surfaces, glass-like input panels (#0F1126), and vibrant neon accents (Cyber Cyan, Electric Blue, Matrix Green, Laser Red, Synthwave Pink).')
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Animations & Interactivity: ').bold = True
    p2.add_run('Features hover slide-in animations for feature cards, dynamic neon glow transitions on buttons, and gradient line accents.')
    
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('Advanced Camera HUD: ').bold = True
    p3.add_run('The real-time camera scanning window is styled as a futuristic heads-up display. It features a sweeping scanner line, dynamic tech-aesthetic corner crosshairs, Viola-Jones global search grid visualizations, color-coded bounding boxes (Neon Green for known faces, Orange/Red for analyzing), and real-time scanning percentage bars.')
    
    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run('Audio Feedback: ').bold = True
    p4.add_run('Integrated pyttsx3 Text-to-Speech (TTS) engine and dynamic micro-sounds (frequency beeps) provide auditory confirmations during registration and scanning, enhancing the premium feel.')
    
    # Section 5
    doc.add_heading('5. SYSTEM FEATURES', level=1)
    
    p1 = doc.add_paragraph(style='List Number')
    p1.add_run('Student Registration: ').bold = True
    p1.add_run('Seamlessly captures 100 face images in real-time and trains the LBPH model with interactive progress tracking.')
    
    p2 = doc.add_paragraph(style='List Number')
    p2.add_run('Subject Registration: ').bold = True
    p2.add_run('Manages classes, faculty names, and distinct time slots.')
    
    p3 = doc.add_paragraph(style='List Number')
    p3.add_run('Live Attendance Scanning: ').bold = True
    p3.add_run('Scans for a predefined interval (20 seconds), tracking and identifying matching faces. It strictly eliminates duplicate entries in real-time.')
    
    p4 = doc.add_paragraph(style='List Number')
    p4.add_run('Data Analytics & Records: ').bold = True
    p4.add_run('Automatically generates professional Excel (.xlsx) and CSV reports with timestamps. Users can browse these attendance sheets gracefully via high-contrast UI Treeview tables.')
    
    # Section 6
    doc.add_heading('6. CONCLUSION & FUTURE SCOPE', level=1)
    p = doc.add_paragraph('NexAttend successfully bridges advanced computer vision algorithms with a highly polished, interactive frontend. By visually demonstrating the Viola-Jones object detection algorithm within a premium framework, it serves as an outstanding capstone project that meets and exceeds enterprise-level visual standards. Future improvements may include integrating cloud-based attendance synchronization and providing a cross-platform mobile application for real-time administrator viewing.')
    
    # Save document
    doc.save('D:\\face-attendance-system\\NexAttend_Project_Report.docx')
    print("Document successfully created at D:\\face-attendance-system\\NexAttend_Project_Report.docx")

if __name__ == "__main__":
    create_report()
