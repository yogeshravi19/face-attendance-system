import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

def create_placeholder(filename, title, subtitle):
    # Generates a premium dark-mode placeholder image
    size = (800, 450)
    img = Image.new('RGB', size, color=(5, 5, 17)) # Deep space blue
    draw = ImageDraw.Draw(img)
    # Glowing border
    draw.rectangle([10, 10, size[0]-10, size[1]-10], outline=(37, 99, 235), width=4)
    draw.rectangle([15, 15, size[0]-15, size[1]-15], outline=(0, 240, 255), width=1)
    
    # Text
    try:
        font_large = ImageFont.truetype("arial.ttf", 40)
        font_small = ImageFont.truetype("arial.ttf", 20)
    except:
        font_large = ImageFont.load_default()
        font_small = ImageFont.load_default()
        
    draw.text((size[0]//2 - 200, size[1]//2 - 30), title, fill=(255, 255, 255), font=font_large)
    draw.text((size[0]//2 - 150, size[1]//2 + 20), subtitle, fill=(130, 140, 160), font=font_small)
    img.save(filename)

# Generating Image Assets
create_placeholder("dashboard_ui.png", "NEXATTEND MAIN DASHBOARD", "VisioMark Premium Dark Theme Layout")
create_placeholder("scanning_hud.png", "LIVE SCANNING HUD", "Viola-Jones Detection & Sweeping Scanner")
create_placeholder("excel_output.png", "EXCEL REPORT OUTPUT", "Automated CSV / XLSX Generation")

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h

# --- Title ---
title = doc.add_heading('NEXATTEND: AI-POWERED FACE RECOGNITION ATTENDANCE SYSTEM', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Comprehensive Phase II Project Report\n', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 1. Abstract ---
add_heading('1. ABSTRACT', 1)
doc.add_paragraph(
    "Institutions and organizations worldwide currently face significant challenges with traditional, manual attendance "
    "tracking systems. These legacy methods—such as roll calls or paper-based sign-in sheets—are not only highly time-consuming "
    "but are inherently prone to human error, proxy attendance, and administrative inefficiencies. As educational institutions "
    "and corporate entities scale, the requirement for an automated, highly secure, and contactless biometric verification "
    "system becomes undeniable.\n\n"
    
    "NexAttend is a state-of-the-art, AI-powered face recognition attendance system designed explicitly to eradicate the "
    "shortcomings of manual roll-call paradigms. By seamlessly integrating modern computer vision architectures and machine "
    "learning classifications, NexAttend accomplishes high-speed facial detection and recognition in real time. The proposed "
    "system captures the temporal presence of individuals, accurately resolving identities against a highly trained biometric "
    "database, and autonomously logs the verified attendance into structured tabular formats, specifically CSV and modern "
    "Microsoft Excel (.xlsx) reports.\n\n"
    
    "This report comprehensively details the development, mathematical foundation, architecture, and user interface design of NexAttend. "
    "The underlying computer vision pipeline optimally fuses the Viola-Jones object detection algorithm (leveraging Haar Cascade Classifiers) "
    "for hyper-fast facial localization with the Local Binary Patterns Histograms (LBPH) recognizer for robust, lighting-invariant "
    "identity classification. Beyond the algorithmic backend, NexAttend introduces the 'VisioMark Premium Dark Theme'—a completely "
    "overhauled graphical user interface (GUI) crafted using Python’s Tkinter library. The GUI integrates glassmorphism, dynamic "
    "hover animations, real-time auditory feedback mechanisms, and a futuristic 'heads-up display' (HUD) for live camera scanning, "
    "thereby ensuring that the software perfectly balances industrial-grade accuracy with an elite, enterprise-level user experience."
)

# --- 2. Introduction ---
add_heading('2. INTRODUCTION', 1)
add_heading('2.1 Overview', 2)
doc.add_paragraph(
    "The 21st century has been characterized by the rapid integration of artificial intelligence into routine administrative "
    "tasks. Biometric attendance systems have historically relied on fingerprint scanners or RFID (Radio Frequency Identification) "
    "tags. While these were significant improvements over pen-and-paper, they introduced their own vulnerabilities—RFID cards "
    "could be passed to colleagues to facilitate proxy attendance, and fingerprint scanners posed contact-based hygiene risks, "
    "as made glaringly apparent during global health crises.\n\n"
    
    "Facial recognition technology solves these multifaceted problems by offering a contactless, non-intrusive, and highly "
    "secure method of identity verification. Because a human face cannot be easily forged or shared, identity spoofing is "
    "significantly mitigated. NexAttend capitalizes on these technological advancements by delivering an end-to-end desktop "
    "application designed to register, manage, and track individual attendance purely through optical facial scanning."
)

add_heading('2.2 Objectives', 2)
doc.add_paragraph(
    "The primary objective of this project is to develop a highly responsive, standalone computer vision attendance system "
    "capable of operating in real-time on consumer-grade hardware. The secondary objectives include:\n"
    "1. High-Accuracy Face Detection: Minimizing false positives and false negatives during the scanning phase by utilizing highly optimized Haar Cascades.\n"
    "2. Lighting Resilience: Ensuring the facial recognition engine (LBPH) can identify individuals under diverse internal lighting conditions.\n"
    "3. Premium User Interface: Discarding the outdated, rudimentary graphical interfaces commonly associated with academic python projects in favor of a sleek, dark-themed, glass-morphism aesthetic.\n"
    "4. Zero-Friction Reporting: Automating the administrative backend heavily so that raw recognition data is instantly and beautifully compiled into readable spreadsheets.\n"
    "5. Scalability: Allowing the system to register hundreds of individuals and map them to distinctive organizational subjects, time slots, and faculty members."
)

# --- 3. Literature Review ---
add_heading('3. LITERATURE REVIEW & THEORETICAL FOUNDATION', 1)
doc.add_paragraph(
    "Understanding the mechanics behind NexAttend requires a foundational exploration of the two vital computer vision algorithms "
    "that drive the system: The Viola-Jones framework and Local Binary Patterns Histograms (LBPH)."
)

add_heading('3.1 The Viola-Jones Object Detection Algorithm', 2)
doc.add_paragraph(
    "Introduced by Paul Viola and Michael Jones in 2001, this framework represented a monumental leap in the capabilities of computer "
    "vision. At a time when image processing was severely constrained by hardware limitations, Viola-Jones offered an algorithm capable "
    "of detecting faces securely at high frame rates. It relies on four core concepts: Haar-like Features, the Integral Image, the "
    "AdaBoost learning algorithm, and the Cascade Classifier.\n\n"
    
    "Haar-like features analyze the varying intensities of pixels in a localized quadrant. For example, in a human face, the region of "
    "the eyes is typically significantly darker than the region of the upper cheeks. The Viola-Jones algorithm scans the screen searching "
    "for these specific contrast arrangements. To compute these rectangular feature matrices in O(1) time complexity, it utilizes "
    "an 'Integral Image' representation. Following this, AdaBoost is deployed to select only the most critical features from "
    "hundreds of thousands of possibilities. Finally, to optimize performance, the algorithm employs a 'Cascade Classifier'—a "
    "series of increasingly complex staging filters. If a sub-window of an image fails the first, simplest filter, it is instantly "
    "discarded, saving immense computational power. In NexAttend, we explicitly utilize the 'haarcascade_frontalface_default.xml' "
    "to enforce this high-speed, lightweight detection within our dynamic HUD."
)

add_heading('3.2 Local Binary Patterns Histograms (LBPH) Recognizer', 2)
doc.add_paragraph(
    "Once the Viola-Jones algorithm isolates the bounding box of a human face, the cropped region of interest (ROI) is passed to the "
    "LBPH recognizer. While basic approaches like Eigenfaces or Fisherfaces consider the entire image grid fundamentally, LBPH focuses "
    "on extracting the localized texture and micro-patterns of a face.\n\n"
    
    "The underlying mathematical procedure iterates through every pixel in the grayscale cropped face. For a given pixel, it examines "
    "its 8 immediate neighbors. If a neighbor’s intensity is greater than or equal to the center pixel, it is assigned a '1'. If lower, "
    "it is assigned a '0'. Reading these binary digits sequentially yields an 8-bit string, which is then converted into a decimal value. "
    "This calculation creates a new image matrix completely resistant to monolithic lighting changes since the absolute light intensity "
    "is ignored in favor of relative neighboring intensity. Finally, the processed face is divided into a geometric grid, and a standard "
    "histogram is extracted from each sector. By concatenating these histograms, the system forms a highly characteristic geometrical "
    "fingerprint of the identity. NexAttend utilizes OpenCV’s 'cv2.face.LBPHFaceRecognizer_create()' to calculate the Euclidean distance "
    "between the real-time histogram and the 100 trained histograms per student, yielding a mathematically concrete confidence metric."
)

# --- 4. System Architecture ---
add_heading('4. SYSTEM ARCHITECTURE & METHODOLOGY', 1)
doc.add_paragraph(
    "The NexAttend architecture is structured as a monolithic standalone desktop application written essentially in Python 3. "
    "It separates logic into clear, cohesive subsystems: User Interface, Data Management, Training Pipeline, and Real-Time Recognition."
)

add_heading('4.1 System Workflow', 2)
doc.add_paragraph(
    "The user initially interacts with the system through the primary Main Dashboard. A fresh installation requires the administrator "
    "to select the 'Register Student' protocol. The user inputs their unique enrollment identification and full name. NexAttend initializes "
    "the local web camera and executes a high-speed data acquisition loop, successfully capturing 100 grayscale, normalized images of the "
    "user's face. These files are securely saved to local storage within the 'TrainingImage' directory.\n\n"
    
    "Subsequent to acquisition, the administrator clicks 'Train Model'. The system traverses the dataset, translating the raw optical "
    "data into LBPH histogram datasets, saving the compiled mathematical model as 'Trainner.yml'. Concurrently, user metadata is appended "
    "to a central CSV database. The system is then prepared for the operational phase. During operation, an administrator selects a "
    "registered subject entity (e.g., 'Computer Networks - Slot A1') from a tabulated list and activates the scanner. The software binds "
    "to the camera and performs live matching against the 'Trainner.yml' state machine."
)

# --- 5. UI/UX Design ---
add_heading('5. UI/UX DESIGN: THE VISIOMARK THEME', 1)
doc.add_paragraph(
    "Standard academic implementations using Python Tkinter are notoriously characterized by gray, archaic Windows 95 aesthetics. "
    "A paramount objective of NexAttend was to completely overthrow this paradigm by implementing 'VisioMark'—a bespoke, premium dark theme."
)

# Insert Dashboard Image
doc.add_paragraph("[ Figure 1: NexAttend Main Dashboard Interface ]").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture("dashboard_ui.png", width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('5.1 Aesthetic Principles', 2)
doc.add_paragraph(
    "The VisioMark system employs a deeply saturated space-black background (#050511) augmented with elevated card-based architectures "
    "(#0D0E20). To simulate depth, 'Glassmorphism' principles are applied to input fields using translucent-like navy blue layers "
    "(#0F1126). Critical action areas are intensely highlighted using a neon color palette comprising Cyber Cyan, Electric Blue, "
    "Matrix Green, and Laser Red. The interface is meticulously constructed using Canvas polygons to guarantee perfectly smooth, "
    "anti-aliased rounded corners on all cards and interactive buttons, a feature natively lacking in standard Tkinter.\n\n"
    
    "Dynamic animation heavily drives user engagement. When the application launches, the core feature cards individually traverse "
    "into view utilizing a carefully timed Y-axis coordinate translation loop (slide_up animation). When the user hovers interactively "
    "over these cards, the borders glow brilliantly using synthetic drop-shadow lines, and high-frequency 'micro-sounds' are emitted "
    "using the Windows Winsound API to provide absolute tactile feedback."
)

add_heading('5.2 The Advanced Camera HUD', 2)
doc.add_paragraph(
    "During the 'Take Attendance' phase, the camera interface does not merely display a raw video feed. Instead, it overlays "
    "a highly technical, cinematic 'Heads-Up Display' (HUD) to visually interpret the complexities of the underlying algorithms."
)

# Insert HUD Image
doc.add_paragraph("[ Figure 2: Live Scanning Heads-Up Display ]").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture("scanning_hud.png", width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "The HUD implements a Viola-Jones Global Search Grid, drawing a faint scanning lattice over the entire screen to represent "
    "the window-sliding detection mechanism. Once a face is detected, standard bounding boxes are replaced with tech-aesthetic corner "
    "crosshairs. Furthermore, as an educational visualization, the system explicitly draws pink and yellow bounding boxes around the "
    "suspected 'Eye Region' and 'Nose/Cheek Region'—the exact Haar-like features the algorithm searches for.\n\n"
    
    "To display confidence rates intuitively, the mathematical Euclidean distance is inverted and normalized into a clean 0-100% metric. "
    "A known user triggers a Matrix Green overlay, while unrecognized entities trigger a warning Orange/Red state. A dynamic 20-second "
    "progress bar shrinks linearly beneath the user's face, dictating the duration of the scan sequence."
)

# --- 6. Results and Output ---
add_heading('6. RESULTS, ANALYTICS & OUTPUT', 1)
doc.add_paragraph(
    "Upon the termination of the scanning protocol (typically after 20 seconds of operation), NexAttend intelligently halts the camera "
    "subroutine and instantiates the data compilation sequence. The Pandas library heavily manipulates the captured identification IDs. "
    "Duplicate recognitions—a common issue where the system recognizes the identical individual 50 times in 10 seconds—are swiftly purged "
    "using DataFrame sanitization techniques. The remaining absolute identifications are stamped with real-world cryptographic time data."
)

# Insert Excel Image
doc.add_paragraph("[ Figure 3: Automated Excel Export resulting from a Session ]").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture("excel_output.png", width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "The system autonomously provisions uniquely named CSV and Microsoft Excel (.xlsx) files formatted identically with headers specifying "
    "Student Name, Enrollment Integer, and Status ('Present'). These documents are natively stored under specifically organized hierarchical "
    "directories aligned with their registered academic subjects (e.g., 'Attendance/MachineLearning/ML_2026-04-10_14-30-00.xlsx'). "
    "Additionally, the users can natively query and view these records directly from the main application via an interlinked Treeview grid, "
    "effectively mitigating the prerequisite of an external spreadsheet viewer for simple checks."
)

# --- 7. Challenges and Optimization ---
add_heading('7. CHALLENGES & SYSTEM OPTIMIZATION', 1)
doc.add_paragraph(
    "Throughout the development lifecycle, several structural challenges emerged. Firstly, performing high-resolution optical captures "
    "and OpenCV analytics on the main sequential thread of a graphical user interface instigates severe freezing and blocking of the UI "
    "canvas. To resolve this, the system’s camera read functions are intelligently offloaded and cycled using the non-blocking `win.after()` "
    "Mainloop callbacks, guaranteeing a consistent 60 Frames Per Second UI fluidity whilst simultaneously digesting structural algorithms.\n\n"
    "Secondly, variations in lighting posed detection failures. This was aggressively mitigated by coercing all camera frames into a strict "
    "grayscale matrix, followed inherently by standard Histogram Equalization. `cv2.equalizeHist()` aggressively stretches the contrast bounds "
    "of the image, resolving issues where users were masked by harsh backlights."
)

# --- 8. Conclusion ---
add_heading('8. CONCLUSION & FUTURE SCOPE', 1)
doc.add_paragraph(
    "NexAttend manifests an authoritative intersection between advanced artificial intelligence processing and spectacular frontend implementation. "
    "By bridging Local Binary Pattern Histograms with Haar Cascades, the software unequivocally achieves an enterprise-grade standard for identity "
    "tracking. The introduction of the VisioMark theme explicitly dismisses the stereotype that academic python frameworks cannot possess "
    "market-level visual prowess.\n\n"
    
    "Future iterations of NexAttend will prioritize cloud interoperability, enabling synchronicity of attendance metrics to distant SQL or Firebase "
    "real-time databases. Moreover, adopting Deep Learning pipelines (such as ResNet-50 implementations via frameworks like PyTorch or TensorFlow) "
    "could replace LBPH, introducing unparalleled multi-angle face tracking capabilities. Ultimately, NexAttend successfully fulfills its objective, "
    "delivering a robust, sophisticated, and autonomous platform for modern organizational infrastructure."
)

# Save Document
try:
    doc.save("D:/face-attendance-system/NexAttend_Detailed_Report.docx")
    print("Success")
except Exception as e:
    print(f"Error: {str(e)}")
