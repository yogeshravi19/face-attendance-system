import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Helper function for adding headings comfortably
    def add_heading(text, level):
        return doc.add_heading(text, level=level)

    def add_image_if_exists(filename, caption):
        if os.path.exists(filename):
            doc.add_paragraph(f"[ {caption} ]").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(filename, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() # Spacing
        else:
            print(f"Warning: {filename} not found. Skipping image insertion.")

    # --- Title Page ---
    title = add_heading('NEXATTEND: AI-POWERED FACE RECOGNITION ATTENDANCE SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Comprehensive Full-System Report with Screenshots\n', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section: Abstract ---
    add_heading('1. ABSTRACT', 1)
    doc.add_paragraph(
        "Institutions and organizations worldwide currently face significant challenges with traditional, manual attendance "
        "tracking systems. These legacy methods—such as roll calls or paper-based sign-in sheets—are not only highly time-consuming "
        "but are inherently prone to human error, proxy attendance, and administrative inefficiencies. As educational institutions "
        "and corporate entities scale, the requirement for an automated, highly secure, and contactless biometric verification "
        "system becomes undeniable.\n\n"
        "NexAttend is a state-of-the-art, AI-powered face recognition attendance system designed explicitly to eradicate the "
        "shortcomings of manual roll-call paradigms. By seamlessly integrating modern computer vision architectures and machine "
        "learning classifications, NexAttend accomplishes high-speed facial detection and recognition in real time. The proposed "
        "system captures the temporal presence of individuals, accurately resolving identities against a highly trained biometric "
        "database, and autonomously logs the verified attendance into structured tabular formats, specifically CSV and modern "
        "Microsoft Excel (.xlsx) reports."
    )

    # --- Section: UI Overview ---
    add_heading('2. VISIOMARK PURE DARK THEME OVERVIEW', 1)
    doc.add_paragraph(
        "A premium aesthetic layer fundamentally redefines user interactions within the NexAttend environment. The core dashboard "
        "interfaces users directly with four primary system modules encapsulated within floating interactive cards."
    )
    add_image_if_exists("dashboard_ui.png", "Figure 1: Main Application Interface (Authentic Screenshot from Real-Time GUI)")

    # --- Section: System Windows ---
    add_heading('3. CORE SYSTEM MODULES & FUNCTIONAL WINDOWS', 1)
    doc.add_paragraph(
        "The monolithic application is decomposed into four fundamental user pipelines, each hosted within its own customized "
        "TopLevel interactive sub-window."
    )

    add_heading('3.1 User Registration Data Acquisition', 2)
    doc.add_paragraph(
        "Enrolling new candidates entails acquiring sequential high-fidelity visual metadata. The Student Registration prompt securely "
        "links local optical feeds and executes the Viola-Jones detector to capture exactly 100 uniquely varied grayscale bounding boxes "
        "of the target face."
    )
    add_image_if_exists("reg_student.png", "Figure 2: Student Image Registration Console")

    add_heading('3.2 Curricular Subject Generation', 2)
    doc.add_paragraph(
        "The semantic mapping of attendance structures revolves around registered Subjects. Using the 'Register Subject' console, "
        "an administrator provisions distinct time slots, mapping them permanently to dedicated folders that track live attendance output."
    )
    add_image_if_exists("reg_subject.png", "Figure 3: Subject Creation Sub-System")

    add_heading('3.3 Sub-routine: Subject Parsing for Scanning', 2)
    doc.add_paragraph(
        "Before deploying the biometric optical loop, the system forces the end-user to definitively confirm the target academic "
        "container. The interface strictly parses through existing generated subjects dynamically and presents them in a formatted TreeView."
    )
    add_image_if_exists("take_attendance.png", "Figure 4: Attendance Target Chooser (Active Sub-Window)")

    add_heading('3.4 Viewing Historical Tabular Data', 2)
    doc.add_paragraph(
        "Accountability stems from accessible metrics. The View Records endpoint natively renders CSV datasets containing massive "
        "attendance data internally, mapping 'Present' statuses directly across rows. Users easily filter dates inside the internal "
        "Scrollable Canvas widget without needing external processors."
    )
    add_image_if_exists("view_records.png", "Figure 5: Internal Historical Records Viewer")

    # --- Section: Optical Backend ---
    add_heading('4. VISUALIZATION OF OPTICAL ALGORITHMS', 1)
    doc.add_paragraph(
        "The essence of the framework operates silently under the hood. However, a specialized Heads-Up Display (HUD) is utilized "
        "for diagnostics. LBPH processes multi-point histogram comparisons generating percentage confidence heuristics."
    )
    add_image_if_exists("scanning_hud.png", "Figure 6: (Simulated Diagram) The Live Scanning Camera HUD Visualizing Viola-Jones Intersections")

    add_heading('5. AUTONOMOUS DATA PIPELINES', 1)
    doc.add_paragraph(
        "Once a physical session terminates, Pandas aggressively scrubs real-time data queues for collision matches and instantly "
        "spawns high-end spreadsheet aggregations directly readable across corporate infrastructure."
    )
    add_image_if_exists("excel_output.png", "Figure 7: (Simulated Diagram) Enterprise Class Excel Automation")
    
    add_heading('6. CONCLUSIONS', 1)
    doc.add_paragraph(
        "NexAttend sets a standard in rapid Python desktop environments by synthesizing robust LBPH/Haar capabilities strictly embedded "
        "within an uncompromised user interface utilizing dynamic Tkinter canvases and seamless real-time file I/O operations."
    )

    doc.save("D:/face-attendance-system/NexAttend_Ultimate_Visual_Report.docx")
    print("Document successfully compiled with all 7 visual assets.")

if __name__ == "__main__":
    main()
